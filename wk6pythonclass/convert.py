from docx import Document

# Read the content from the file
with open("wk6pythonclass/sample.txt", 'r') as readerobject:
    content = readerobject.read()
     
# Create a new Document
document = Document()

# Add a paragraph with the content from the file
p = document.add_paragraph()

# Add a bold run to the paragraph
paragraph = p.add_run(content)
paragraph.bold = True
paragraph.italic = True


# Add an italic run to the paragraph


records = (
    (1, '102004', 'Gbolahan gboy'),
    (2, '102005', 'Thomas Joshua'),
    (3, '102006', 'Thomoty mark')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'ID'
hdr_cells[1].text = 'schoolNO'
hdr_cells[2].text = 'Name'
for ID,schoolno , Names in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(ID)
    row_cells[1].text = schoolno
    row_cells[2].text = Names

document.add_page_break()

document.save('students.docx')

# #open an existinfg word 


