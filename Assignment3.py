'''
Working with Word documents in python (Document Creation and Basic Operations)

'''
# ● Create a new Word document.
# ● Add a title with a specific style (e.g., centered, bold, and larger font size).
# ● Save the document.
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_TAB_ALIGNMENT
from docx.shared import Pt, RGBColor
from docx.enum.section import WD_SECTION


document = Document()
p = document.add_heading(level=0)
title = p.add_run('TRIBUTE TO ALBERT EINSTEIN')
title.bold = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER # Align to the center:
title.font.size = Pt(34)  # larger fontsize: setting font size to 24 points


# Add a paragraph with multiple runs (each run having different styles such as bold, italic, and underline).
intro = document.add_paragraph()
intro1 = intro.add_run("Albert Einstein ")
intro1.font.color.rgb = RGBColor(39, 48, 39) # addding font color to the text
intro1.italic = True # adding italize font styling
intro1.bold = True
intro1.font.size = Pt(16) 

intro2 = intro.add_run(" was born in 1879 ")
intro2.bold= True
intro2.font.color.rgb = RGBColor(255, 0, 0)  # Set the color to red
intro3 = intro.add_run(" in Germany, the first child of a bourgeois Jewish couple. ")


intro4 = intro.add_run("The young Albert displayed an early interest in science, ")
intro4.underline = True
intro5 = intro.add_run(
    "but he was unhappy with the principles of obedience and conformity that governed his Catholic elementary school.")


#  Create a bulleted list with at least three items, each with different font styles
bullet_point_1 = document.add_paragraph(style='ListBullet')
txt1 = bullet_point_1.add_run("Developed the theory of relativity")
txt1.bold = True

bulletpoint2 = document.add_paragraph(style='ListBullet')
txt2 = bulletpoint2.add_run("Won the Nobel Prize in Physics in 1921")
txt2.italic = True

bulletpoint3 = document.add_paragraph(style='ListBullet')
txt3 = bulletpoint3.add_run("Made significant contributions to the philosophy of science")
txt3.font.color.rgb = RGBColor(0, 0, 255)  # Set the color to blue



img = document.add_picture('Einstein.jpg', width=Inches(5.45), height=Inches(3.00))

# Center-align the image
last_paragraph = document.paragraphs[-1]
last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

records = (
    ('Theory of General Relativity', '1915','Fundamental theory in physics explaining gravitation '),
    ('Photoelectric Effect', '1905','Discovery explaining interaction of light with matter'),
    ('Brownian Motion', '1905','Experimental evidence supporting atomic theory'),
   ( 'Mass-Energy Equivalence (E=mc^2)', '1905','Famous equation demonstrating equivalence of mass and energy')
)

table = document.add_table(rows=1, cols=3)
# Set table style and alignment
table.style = 'Table Grid'
table.alignment = WD_TAB_ALIGNMENT.CENTER

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Invention Name'
hdr_cells[1].text = 'Year'
hdr_cells[2].text = 'Description'
for invent, year, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = invent
    row_cells[1].text = str(year)
    row_cells[2].text = desc


#  Sections and Headers/Footers
# ● Create multiple sections in the document.
# ● Add different headers and footers for each section.
# ● Include page numbers in the footers.
 
nextpage = document.add_page_break()
# d section 1 with a different header and footer
section1 = document.sections[0]
section1.start_type
section1 = document.add_section(WD_SECTION.NEW_PAGE)
header1 = section1.header
header1_paragraph = header1.paragraphs[0].add_run('Header ')
header1_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

footer1 = section1.footer
footer1_paragraph = footer1.paragraphs[0].add_run('''Everybody is a genius. But if you judge a fish by its ability to climb a tree, it will live its whole life believing that it is stupid"
~~ Albert Einstein''')

bulletpoint4 = document.add_paragraph(style='ListBullet')
txt4 = bulletpoint4.add_run('''He is best known to the general public for his mass–energy equivalence formula E = mc2, which has been dubbed "the world's most famous equation".
                            ''')
txt4.font.color.rgb = RGBColor(85, 99, 96)  # Set the color 

bulletpoint5 = document.add_paragraph(style='ListBullet')
txt5 = bulletpoint5.add_run('''Einstein published more than 300 scientific papers and more than 150 non-scientific works. His intellectual achievements and originality have made the word "Einstein" synonymous with "genius".  ''')
txt5.font.color.rgb = RGBColor(85, 99, 96)  # Set the color 

bulletpoint6 = document.add_paragraph(style='ListBullet')
txt6 = bulletpoint6.add_run('''He received the 1921 Nobel Prize in Physics "for his services to theoretical physics, and especially for his discovery of the law of the photoelectric effect", a pivotal step in the development of quantum theory.
                              ''')
txt6.font.color.rgb = RGBColor(85, 99, 96)  # Set the color 

bulletpoint7 = document.add_paragraph(style='ListBullet')
txt7 = bulletpoint7.add_run('''He lived in Switzerland between 1895 and 1914, except for one year in Prague, and he received his academic diploma from the Swiss federal polytechnic school (later the Eidgenössische Technische Hochschule, ETH) in Zürich in 1900. He taught theoretical physics there between 1912 and 1914 before he left for Berlin. ''')
txt7.font.color.rgb = RGBColor(85, 99, 96)  # Set the color





document.save("TributeEinstein.docx")