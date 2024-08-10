from docx import Document  #This is a python external library that allows us to read .docx files
import openpyxl   # This is a python external library  that allows us to create and write to Excel files
import os  # This library allows us to interact with the file system


#  this i s a Function to extract data from a .docx file  on the same directory
def extract_invoice_data(docx_file):
    doc = Document(docx_file)
    invoice_data = {
        "Invoice ID": "",
        "Total Products": 0,
        "Subtotal": 0.0,
        "Tax": 0.0,
        "Total": 0.0
    }
    
    # Extract Invoice ID
    invoice_data["Invoice ID"] = doc.paragraphs[0].text.replace("INV", "")
    
    # Extract product information
    product_count = 0
    for para in doc.paragraphs[1:]:
        if para.text.startswith("SUBTOTAL"):
            break
        if ":" in para.text:
            parts = para.text.split(":")
            if len(parts) == 2:
                product, qty = parts
                product_count += int(qty)
    
    invoice_data["Total Products"] = product_count
    
    # Extract subtotal, tax, and total
    for para in doc.paragraphs[1:]:
        text = para.text.strip()
        if "SUBTOTAL" in text:
            invoice_data["Subtotal"] = float(text.split(":")[1].split()[0])
        elif "TAX" in text:
            invoice_data["Tax"] = float(text.split(":")[1].split()[0])
        elif "TOTAL" in text:
            invoice_data["Total"] = float(text.split(":")[1].split()[0])
    
    return invoice_data

# Function to create a spreadsheet from the extracted data
def create_spreadsheet(invoice_data_list, output_file):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Invoices"
    
    # Create headers
    headers = ["Invoice ID", "Total Products", "Subtotal", "Tax", "Total"]
    ws.append(headers)
    
    # to  Add data s
    for invoice_data in invoice_data_list:
        ws.append([
            invoice_data["Invoice ID"],
            invoice_data["Total Products"],
            invoice_data["Subtotal"],
            invoice_data["Tax"],
            invoice_data["Total"]
        ])
    
    # Save the workbook
    wb.save(output_file)

# Main function to process all invoices
def process_invoices(docx_folder, output_file):
    invoice_data_list = []
    
    # Process each .docx file in the folder
    for filename in os.listdir(docx_folder):
        if filename.endswith(".docx") and not filename.startswith("~$"): 
            filepath = os.path.join(docx_folder, filename)
            invoice_data = extract_invoice_data(filepath)
            invoice_data_list.append(invoice_data)
    
    # Create the spreadsheet
    create_spreadsheet(invoice_data_list, output_file)

# Specify the folder containing .docx files and the output Excel file
docx_folder = "Openpyl_file"  #  folder containing the .docx files
output_file = "output_invoice.xlsx"    #this will be Output Excel file name

# call the function
process_invoices(docx_folder, output_file)
