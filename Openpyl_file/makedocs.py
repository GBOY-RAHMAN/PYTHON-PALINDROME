from docx import Document
import random
import math

def makeInvoices(numFiles):
    products = ["Parka", "Boots", "Snowshoes", "Climbing Rope", "Oxygen Tank", "Ice Pick", "Crampons"]

    # Invoice loop
    for i in range(numFiles):
        
        # Create Randomized invoice
        invoiceNum = "100" + str(i).zfill(4)
        productList = {}
        for j in range(random.randint(1,10)):
            product = products[random.randint(0,len(products)-1)]
            if product in productList:
                productList[product] += 1
            else:
                productList[product] = 1
        subtot = round(random.random()*10**(random.randint(3, 4)), 2)
        tax = round(subtot*0.13, 2)
        total = round(subtot + tax, 2)

        # Create doc from random invoice
        aDoc = Document()
        aDoc.add_heading("INV" + invoiceNum)
        pProd = aDoc.add_paragraph("PRODUCTS\n")
        for key in productList.keys():
            pProd.add_run(f"{key}:{productList[key]}\n")
        aDoc.add_paragraph(f"SUBTOTAL:{subtot}\nTAX:{tax}\nTOTAL:{total}")
        aDoc.save(f"INV{invoiceNum}.docx")


makeInvoices(5)